from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading('RAG Chatbot Evaluator MVP - Detailed Report', 0)

doc.add_heading('Overview', level=1)
doc.add_paragraph('This MVP evaluates RAG chatbot answers against ground-truth context using a judge LLM. It scores correctness, completeness, and relevance on a 1-5 scale, provides reasoning, and aggregates results into a failure mode report.')

doc.add_heading('Implementation Details', level=1)
doc.add_paragraph('''
- Language: Python
- Libraries: openai (for LLM API), python-docx (for this report)
- Input: JSON file with mock data (questions, RAG answers, ground-truth)
- Output: evaluations.json (detailed scores/reasoning), failure_report.json (aggregated low scores)
- LLM: OpenAI GPT-4 (or mock mode for testing)
- No front-end; command-line script
''')

doc.add_heading('Mock Data', level=1)
doc.add_paragraph('Sample data includes 3 questions with varying answer quality:')
doc.add_paragraph('1. Capital of France - Correct answer')
doc.add_paragraph('2. Author of Pride and Prejudice - Correct answer')
doc.add_paragraph('3. Boiling point of water - Incorrect answer (90°C instead of 100°C)')

doc.add_heading('Evaluation Process', level=1)
doc.add_paragraph('''
For each RAG answer:
- Prompt LLM to score correctness, completeness, relevance (1-5)
- Provide reasoning for each score
- Aggregate low scores (<3) into failure report
''')

doc.add_heading('Testing Results', level=1)
doc.add_paragraph('''
- Happy path: Script loads data, evaluates answers, generates reports
- Error handling: Mock mode for missing API key
- Output verification: Scores and reasoning match expected values
- Failure report: Correctly identifies low-scoring answers
''')

doc.add_heading('Example Outputs', level=1)
doc.add_paragraph('Evaluations JSON includes scores like:')
doc.add_paragraph('Correctness: 5 (accurate), Completeness: 4 (lacks context), Relevance: 5 (direct)')
doc.add_paragraph('Failure report highlights issues: e.g., "correctness: Incorrect boiling point"')

doc.add_heading('Setup and Usage', level=1)
doc.add_paragraph('''
1. Install dependencies: pip install -r requirements.txt
2. Set OPENAI_API_KEY (optional for mock mode)
3. Run: python evaluator.py
4. Check output files: evaluations.json, failure_report.json
''')

doc.add_heading('Limitations and Extensions', level=1)
doc.add_paragraph('''
- Limited to 3 samples; can scale to more data
- Uses single LLM; could add multiple judges
- No UI; could add web interface
- Simple aggregation; could add charts/visualizations
- Mock mode for testing; real API for production
''')

doc.save('rag_evaluator_detailed_report.docx')
print("Detailed report saved as rag_evaluator_detailed_report.docx")
